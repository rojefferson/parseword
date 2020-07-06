import os
import docx
from docx.oxml.ns import qn
import mammoth
from bs4 import BeautifulSoup
from Conexao import Conexao
from Venda import Venda
import re

def findChecked(text):
    retorno = []
    fd = text.find("<w:checked")
    while fd > 0:
      text = text[fd:]
      value = text.find(">")
      tag  = text[0:value]
      if tag.find("0")> 0:
        retorno.append("0")
      else:
        retorno.append("1")
      text = text[value:]
      fd = text.find("<w:checked")
    return retorno

  
def checkboxInsert(table):        
  listCheckbox = list([("categria",[]) ,("Pagamento",[]),("BotaoExtra",[])])
  for i, row in enumerate(table.rows):
    for cell in row.cells:
      for paragraph in cell.paragraphs:
        p = paragraph._element
        checkBoxes = p.xpath('.//w:checkBox')
        if(len(checkBoxes) == 4 and len(listCheckbox[0][1]) == 0):
          listCheckbox[0] = tuple(("categria",findChecked(p.xml)))
        elif(len(checkBoxes) == 5 and len(listCheckbox[1][1]) == 0):
          listCheckbox[1] = tuple(("pagamento",findChecked(p.xml)))
        elif(len(checkBoxes) == 2 and len(listCheckbox[2][1]) == 0):
          listCheckbox[2] = tuple(("BotaoExtra",findChecked(p.xml)))
  return listCheckbox   

def removeTagStrong(tag):
    [x.extract() for x in tag.findAll('strong')]
    return tag.text

def trataDataText(dataText):
  dataArray = re.findall(r'\d+',dataText)
  dataArray = list(reversed(dataArray))
  retorno = "-".join(dataArray)
  return retorno

def getVendaDoc(fullpath,table):
   with open(fullpath, "rb") as docx_file:
      result = mammoth.convert_to_html(docx_file)
      html = result.value
      soup = BeautifulSoup(html)
      alltr = soup.findAll("tr")   
      contrato = removeTagStrong(alltr[0].find("td"))
      chip = removeTagStrong(alltr[0].findAll("td")[1])
      rastreador = removeTagStrong(alltr[1].findAll("td")[0])
      linha = removeTagStrong(alltr[1].findAll("td")[1])
      cliente = alltr[3].findAll("p")[1].text
      CPFCNPJ =  alltr[4].findAll("p")[1].text
      RGIEre =   re.findall(r'\d+', alltr[4].findAll("p")[3].text)  
      RGIE = ""
      if (len(RGIEre) > 0):
        RGIE = alltr[4].findAll("p")[3].text
      dataNascimento =  trataDataText(str(alltr[4].findAll("p")[5].text))
      endereco = alltr[5].findAll("p")[1].text
      bairro   = alltr[6].findAll("p")[1].text
      cidadeUF = alltr[6].findAll("p")[3].text
      cep = alltr[6].findAll("p")[5].text
      telefoneResidencial = alltr[7].findAll("p")[1].text
      telefoneCelular = ""
      telefoneCelular =  " - ".join(list(map(lambda x : x.text, alltr[7].findAll("p")[3:])))
      email  =  alltr[8].findAll("p")[1].text
      marca =  alltr[11].findAll("p")[1].text
      modeloAno = alltr[11].findAll("p")[3].text
      cor = alltr[12].findAll("p")[1].text
      placa = alltr[12].findAll("p")[3].text
      panico1 = alltr[13].findAll("p")[1].text
      panico2 = alltr[14].findAll("p")[1].text
      valorRastreadoRE=   re.findall(r'\d+',alltr[16].findAll("p")[1].text)
      Valorbotao = ''
      valorBotaoRE =   re.findall(r'\d+', alltr[16].findAll("p")[3].text)
      if(len(valorBotaoRE) > 0):
        Valorbotao = alltr[16].findAll("p")[3].text
      parcelasRE =  re.findall(r'\d+',alltr[17].findAll("p")[1].text)
      qtdBotaoRE =  re.findall(r'\d+',alltr[17].findAll("p")[3].text)
      parcelas = 0
      qtdBotaoExtra = 0

      valorRastreador = 0
      if (len(valorRastreadoRE) > 0):       
        valorRastreador = valorRastreadoRE[0]
      if(len(parcelasRE)>0):
        parcelas = parcelasRE[0]
      if(len(qtdBotaoRE) > 0):
        qtdBotaoExtra = qtdBotaoRE[0]
      localVenda = alltr[18].findAll("p")[1].text
      vendedorIndicacao = alltr[18].findAll("p")[1].text
      localInstalacao = alltr[19].findAll("p")[1].text
      instalador = alltr[19].findAll("p")[3].text  
      dataVenda = trataDataText(str(alltr[20].findAll("p")[1].text))
      dataInstalacao = trataDataText(str(alltr[20].findAll("p")[3].text))
      observacoeslist =  alltr[21].findAll("p")
      ob = ""
      for x in observacoeslist:
          ob = ob + ". " + x.text
      check = checkboxInsert(table)
      botaoExtra = "1" if check[2][1] == "1" else "0"
      venda = Venda(contrato,chip,rastreador,linha,cliente,CPFCNPJ,RGIE,dataNascimento,endereco,bairro,cidadeUF,cep,telefoneResidencial,telefoneCelular,email,marca,modeloAno,cor,placa,panico1,panico2,valorRastreador,Valorbotao,localVenda,vendedorIndicacao,localInstalacao,instalador,dataVenda,dataInstalacao,ob,"0",tuple(check[1])[1][0],tuple(check[1])[1][1],tuple(check[1])[1][2],tuple(check[1])[1][3],tuple(check[1])[1][4],botaoExtra,qtdBotaoExtra,tuple(check[0])[1][0],tuple(check[0])[1][1],tuple(check[0])[1][2],tuple(check[0])[1][3],parcelas)
      return venda


def main():
  path = "/home/jefferson/Documentos/arquivos"
  arquivos =  os.listdir(path)
  con = Conexao()
  for arquivo in arquivos:
    fullpath = path + "/" +arquivo
    filename, file_extension = os.path.splitext(fullpath)
    if file_extension == ".docx":
      document = docx.Document(fullpath)
      table = document.tables[0]
      venda =  getVendaDoc(fullpath,table)
      con.insertVenda(venda)

main()
    
        

        
        
        